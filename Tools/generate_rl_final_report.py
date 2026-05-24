from pathlib import Path
import json

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "Assets" / "report_assets"
OUT_DOC = ROOT / "Assets" / "RL_Final_Report_ThaiFixed.docx"
SCALARS = json.loads((ASSET_DIR / "scalars.json").read_text(encoding="utf-8"))

FONT = "Tahoma"


def values(run, tag):
    return [(int(item["step"]), float(item["value"])) for item in SCALARS[run].get(tag, [])]


def last(run, tag):
    data = values(run, tag)
    return data[-1][1] if data else None


def fmt(value, digits=2):
    if value is None:
        return "-"
    return f"{value:,.{digits}f}"


def set_run_font(run, size=14, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), FONT)


def set_paragraph_font(paragraph, size=14):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=run.bold)


def add_p(doc, text="", size=14, align=None, bold=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=18 if level == 1 else 15, bold=True)
        run.font.color.rgb = RGBColor(31, 78, 121)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=14)
    p.paragraph_format.space_after = Pt(2)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    add_p(doc, "")
    return table


def add_picture(doc, filename, caption):
    doc.add_picture(str(ASSET_DIR / filename), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, caption, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)


ppo_final = last("ppo", "Environment/Cumulative Reward")
ppo_len = last("ppo", "Environment/Episode Length")
ppo_len_series = values("ppo", "Environment/Episode Length")
ppo_center = last("ppo", "PidTuningAgent/Reward Breakdown/Centered")
ppo_velocity = last("ppo", "PidTuningAgent/Reward Breakdown/LowVelocity")
ppo_failure = last("ppo", "PidTuningAgent/Reward Breakdown/FailurePenalty")
ppo_total = last("ppo", "PidTuningAgent/Reward Breakdown/TotalTracked")
ppo_policy_loss = last("ppo", "Losses/Policy Loss")
ppo_value_loss = last("ppo", "Losses/Value Loss")

ddpg_recent = last("ddpg", "Policy/Recent Mean Reward")
ddpg_std = last("ddpg", "Policy/Recent Reward Std")
ddpg_episode_reward = last("ddpg", "Environment/Cumulative Reward")
ddpg_len = last("ddpg", "Environment/Episode Length")
ddpg_len_series = values("ddpg", "Environment/Episode Length")
ddpg_center = last("ddpg", "DDPGFC350EPidAgent/Reward Breakdown/Centered")
ddpg_velocity = last("ddpg", "DDPGFC350EPidAgent/Reward Breakdown/LowVelocity")
ddpg_failure = last("ddpg", "DDPGFC350EPidAgent/Reward Breakdown/FailurePenalty")
ddpg_total = last("ddpg", "DDPGFC350EPidAgent/Reward Breakdown/TotalTracked")
ddpg_actor_loss = last("ddpg", "Losses/Actor")
ddpg_critic_loss = last("ddpg", "Losses/Critic")

doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

style = doc.styles["Normal"]
style.font.name = FONT
style.font.size = Pt(14)
style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
style._element.rPr.rFonts.set(qn("w:cs"), FONT)

add_p(doc, "รายงานผลการทดลอง Deep Reinforcement Learning สำหรับระบบ Ball Balancer แบบ 3-RRS", 20, WD_ALIGN_PARAGRAPH.CENTER, True)
add_p(doc, "เปรียบเทียบ PPO และ DDPG-FC-350-E-PID จากข้อมูลผลการ train จริงในโปรเจกต์ Unity", 15, WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "แหล่งข้อมูล: results/pid_tuning_multi_12, results/DDPG_14, Assets/Script, Tools/DDPGFC350EPid และไฟล์โครงรายงานเดิม", 12, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

add_h(doc, "บทคัดย่อ")
add_p(
    doc,
    f"โครงงานนี้ศึกษาการใช้ Deep Reinforcement Learning เพื่อปรับค่า PID gains สำหรับระบบลูกบอลทรงตัวบนแพลตฟอร์มขนานแบบ 3-RRS ที่จำลองใน Unity โดย agent ส่งออกค่า Kp, Ki และ Kd แบบต่อเนื่อง แล้วให้ inner PID controller ใช้ค่าเหล่านี้ควบคุมการเอียงของแพลตฟอร์ม ผลการทดลองขั้นสุดท้ายใช้ PPO run_id pid_tuning_multi_12 และ DDPG run_id DDPG_14 ซึ่ง train ถึงประมาณ 3.5 ล้าน environment steps ทั้งคู่ จาก log พบว่า PPO ให้ cumulative reward ปลายทาง {fmt(ppo_final)} และ episode length {fmt(ppo_len, 0)} steps ส่วน DDPG ให้ recent mean reward {fmt(ddpg_recent)} ค่าเบี่ยงเบนมาตรฐาน {fmt(ddpg_std)} และ episode reward ล่าสุด {fmt(ddpg_episode_reward)} โดยทั้งสองวิธีลด failure penalty เหลือ 0 ในช่วงปลายการ train",
)
add_p(
    doc,
    "ผลโดยรวมชี้ว่า PPO มีความเสถียรของ reward สูงกว่าและ convergence เรียบกว่าในสภาพแวดล้อมนี้ ขณะที่ DDPG-FC-350-E-PID สามารถเรียนรู้พฤติกรรมทรงตัวได้หลังปรับ observation เป็น E-PID และใช้ actor/critic ขนาด 2 hidden layers ชั้นละ 350 nodes แต่ยังมี reward variance สูงกว่า จึงควรพัฒนาต่อด้าน exploration schedule และ critic stability",
)

add_h(doc, "1. วัตถุประสงค์")
for item in [
    "พัฒนาสภาพแวดล้อมจำลองใน Unity สำหรับระบบ ball-balancing บนแพลตฟอร์มขนานแบบ 3-DOF 3-RRS",
    "ออกแบบ agent ให้ปรับค่า PID gains แบบต่อเนื่อง ได้แก่ Kp, Ki และ Kd แทนการส่งคำสั่งมอเตอร์โดยตรง",
    "เปรียบเทียบผลการเรียนรู้ของ PPO และ DDPG-FC-350-E-PID จาก training logs จริง",
    "วิเคราะห์ reward curve, episode length, reward breakdown, hyperparameters และข้อจำกัดของแต่ละ algorithm",
]:
    add_bullet(doc, item)

add_h(doc, "2. พื้นฐานและงานที่เกี่ยวข้อง")
add_p(
    doc,
    "งานอ้างอิงเรื่อง Heuristic and deep reinforcement learning-based PID control of trajectory tracking in a ball-and-plate system ใช้แนวคิดให้ DRL agent ปรับ PID gains สำหรับระบบ ball-and-plate โดยรูปแบบ DDPG-FC-350-E-PID ใช้ actor และ critic ที่มี hidden layer แบบ fully connected จำนวนสองชั้น ชั้นละ 350 nodes และใช้ ELU activation แนวคิดนี้ถูกนำมาปรับใช้กับระบบ 3-RRS ใน Unity ซึ่งมีพลวัตไม่เชิงเส้นจาก rigid-body physics, contact และ friction",
)
add_p(
    doc,
    "PPO เป็น on-policy actor-critic ที่ใช้ clipped surrogate objective จึงคุมขนาดการ update policy ได้ดีและมักเสถียรใน environment ที่ reward มีความผันผวน ส่วน DDPG เป็น off-policy deterministic actor-critic ที่เหมาะกับ action ต่อเนื่องและใช้ replay buffer เพื่อ reuse ประสบการณ์ แต่ไวต่อ hyperparameters, exploration noise และความแม่นของ critic มากกว่า",
)

add_table(
    doc,
    ["ประเด็น", "PPO", "DDPG-FC-350-E-PID"],
    [
        ["Run สุดท้าย", "pid_tuning_multi_12", "DDPG_14"],
        ["Behavior name", "pid_tuning_agent", "ddpg_fc_350_e_pid_agent"],
        ["Observation", "13 ค่า: position, velocity, error, plate angle, PID gains, radial distance, speed", "10 ค่า: error, integral error, derivative error, PID gains, error magnitude โดย error มาจากตำแหน่งเทียบ target และ derivative error แทน velocity เชิงควบคุม"],
        ["Action", "continuous 3 ค่า: Kp, Ki, Kd", "continuous 3 ค่า: Kp, Ki, Kd"],
        ["Network", "2 layers x 128 hidden units", "2 layers x 350 hidden units, ELU"],
        ["Training steps", "final checkpoint 3,500,018", "3,500,000 environment steps"],
    ],
)

add_h(doc, "3. ภาพรวมระบบและการกำหนดปัญหา")
add_p(
    doc,
    "ระบบจำลองประกอบด้วยลูกบอล rigid body บนแพลตฟอร์ม 3-RRS ซึ่งปรับ roll/pitch ผ่าน controller ภายใน โดย learning agent ทำหน้าที่เลือก PID gains ที่ใช้ร่วมกันบนแกน X และ Z การออกแบบนี้ทำให้ DRL เรียนรู้ gain scheduling ที่ตีความได้ง่ายกว่า raw motor command และยังคงประโยชน์ของ PID controller ในการสร้างคำสั่งควบคุมที่ต่อเนื่อง",
)
add_p(
    doc,
    "ปัญหาถูกกำหนดเป็น Markov Decision Process โดย PPO ใช้ observation ที่มีตำแหน่ง ความเร็ว error และสถานะ PID โดยตรง ส่วน DDPG ใช้รูปแบบ E-PID observation ได้แก่ error, integral error และ derivative error ซึ่ง error คำนวณจากตำแหน่งลูกบอลเทียบ target และ derivative error เป็นอัตราการเปลี่ยน error ที่ทำหน้าที่ใกล้เคียง velocity ในเชิงควบคุม; action เป็นค่า PID gains; transition เกิดจาก Unity physics; reward ให้รางวัลกับการอยู่ใกล้ศูนย์กลางและความเร็วต่ำ พร้อมลงโทษเมื่อออกนอกขอบเขตหรือหล่นต่ำกว่าแพลตฟอร์ม",
)

add_h(doc, "4. Reward Function และ Episode Termination")
add_p(doc, "จาก Assets/Script/PidTuningAgent.cs PPO ใช้ reward ต่อ step เป็นผลรวมของ centered reward และ low-velocity reward:")
add_p(doc, "r_PPO = (1 - d_norm^2) * 0.015 + (1 - v_norm) * 0.003")
add_p(doc, "ถ้าลูกบอลออกนอก failDistance = 0.55 หรือระดับความสูงต่ำกว่าแพลตฟอร์ม 0.3 m จะได้รับ failure penalty = -10 และจบ episode")
add_p(doc, "ส่วน DDPG ใช้ reward คล้ายกัน แต่เพิ่ม center progress reward, outward velocity penalty และ PID effort penalty:")
add_p(doc, "r_DDPG = centered + low_velocity + center_progress - outward_velocity_penalty - pid_effort_penalty")
add_p(doc, "DDPG ใช้ E-PID observation ได้แก่ error, integral error และ derivative error โดยไม่ได้ใส่ raw position/velocity เข้าไปตรง ๆ แต่ error มาจากตำแหน่งลูกบอลเทียบ target และ derivative error มาจาก (error - previousError) / dt จึงเป็นตัวแทนอัตราการเคลื่อนที่ของ error นอกจากนี้ยังใช้ seed gains จาก paper คือ Kp=4.2025, Ki=1.2529, Kd=5.1323")

add_h(doc, "5. Hyperparameters")
add_table(
    doc,
    ["รายการ", "PPO: pid_tuning_multi_12", "DDPG: DDPG_14"],
    [
        ["Batch size", "1024", "256"],
        ["Buffer / Replay size", "3000", "200000"],
        ["Learning rate", "0.003 linear schedule", "actor 0.0001, critic 0.001"],
        ["Gamma", "0.99", "0.99"],
        ["GAE lambda / Tau", "lambda 0.95", "tau 0.005"],
        ["Epoch / Update", "num_epoch 3", "update_after 10000, update_every 2"],
        ["Exploration", "entropy beta 0.001", "noise 0.12 -> min 0.01 over 1,200,000 steps"],
        ["Environment", "5 envs, time_scale 20", "5 envs, time_scale 50, no_graphics true"],
    ],
)

add_h(doc, "6. ผลการทดลอง")
add_p(doc, "ผลที่รายงานมาจาก TensorBoard event files และ training_status.json ของ run สุดท้ายที่ระบุไว้ โดยใช้ PPO = pid_tuning_multi_12 และ DDPG = DDPG_14")
add_table(
    doc,
    ["Metric", "PPO", "DDPG-FC-350-E-PID", "วิเคราะห์"],
    [
        ["Final / recent reward", fmt(ppo_final), fmt(ddpg_recent), f"PPO สูงกว่า recent mean ของ DDPG ประมาณ {fmt(ppo_final - ddpg_recent)} reward"],
        ["Latest episode reward", fmt(ppo_final), fmt(ddpg_episode_reward), "episode ล่าสุดของ DDPG ใกล้ PPO แต่ค่าเฉลี่ยล่าสุดยังต่ำกว่า"],
        ["Episode length ปลายทาง", fmt(ppo_len, 0), fmt(ddpg_len, 0), "ทั้งสองเข้าใกล้ 2000 steps แสดงว่ารักษาลูกบอลได้จนเกือบเต็ม episode"],
        ["จำนวน episode length ที่บันทึก", f"{len(ppo_len_series):,} points", f"{len(ddpg_len_series):,} episodes", "DDPG มี Environment/Episode Length ใน log โดยตรง ไม่ได้ขาดข้อมูล"],
        ["Episode length แรก -> สุดท้าย", f"{fmt(ppo_len_series[0][1], 0)} -> {fmt(ppo_len_series[-1][1], 0)}", f"{fmt(ddpg_len_series[0][1], 0)} -> {fmt(ddpg_len_series[-1][1], 0)}", "DDPG เพิ่มจาก episode สั้นช่วงแรกไปถึงเต็ม 2,000 steps ในช่วงท้าย"],
        ["Failure penalty ปลายทาง", fmt(ppo_failure), fmt(ddpg_failure), "เป็น 0 ทั้งคู่ในช่วงท้าย จึงไม่มี failure ใน episode ที่บันทึกล่าสุด"],
        ["Reward centered", fmt(ppo_center), fmt(ddpg_center), "PPO ได้ centered reward สูงกว่า แปลว่าลูกบอลอยู่ใกล้ศูนย์กลางกว่าโดยเฉลี่ย"],
        ["Reward low velocity", fmt(ppo_velocity), fmt(ddpg_velocity), "ใกล้เคียงกันมาก แสดงว่าทั้งสองลดความเร็วลูกบอลได้ดี"],
        ["Loss ปลายทาง", f"policy {fmt(ppo_policy_loss, 4)}, value {fmt(ppo_value_loss, 6)}", f"actor {fmt(ddpg_actor_loss, 4)}, critic {fmt(ddpg_critic_loss, 6)}", "value/critic loss ต่ำ บ่งชี้ว่าการประมาณค่าเริ่มนิ่งในช่วงท้าย"],
    ],
)

add_picture(doc, "reward_comparison.png", "รูปที่ 1 Reward curve เปรียบเทียบ PPO และ DDPG")
add_p(doc, "กราฟแสดงว่า PPO เพิ่ม reward ได้เร็วและนิ่งกว่า โดยเริ่มจากประมาณ 16 และขึ้นไปที่ประมาณ 179.7 ส่วน DDPG เริ่มจากค่าติดลบเพราะ policy ยังสุ่มและมี failure penalty ก่อนค่อย ๆ เพิ่มจนถึง recent mean ประมาณ 170.2 ช่วงปลายการ train DDPG ยังมี variance สูงกว่า PPO")

add_picture(doc, "episode_length_comparison.png", "รูปที่ 2 Episode length เปรียบเทียบความสามารถในการอยู่รอดของ episode")
add_p(doc, f"กราฟนี้ยืนยันว่า DDPG มีข้อมูล Environment/Episode Length ใน TensorBoard log โดยบันทึกทั้งหมด {len(ddpg_len_series):,} episodes ค่าแรกอยู่ที่ {fmt(ddpg_len_series[0][1], 0)} steps และค่าสุดท้ายอยู่ที่ {fmt(ddpg_len_series[-1][1], 0)} steps ส่วน PPO มีค่า episode length ปลายทาง {fmt(ppo_len, 0)} steps ดังนั้นช่วงท้ายทั้งสองวิธีอยู่ได้เกือบเต็ม episode แต่ DDPG มี episode สั้นมากกว่าในช่วงต้นและกลางการ train")

add_picture(doc, "reward_breakdown_final.png", "รูปที่ 3 Reward breakdown ปลายทาง")
add_p(doc, "PPO มี centered reward สูงกว่า DDPG ขณะที่ low-velocity reward ใกล้เคียงกัน และ failure penalty เป็น 0 ทั้งคู่ ความต่างนี้บอกว่า DDPG สามารถทำให้ลูกบอลไม่หลุดได้ แต่ตำแหน่งเฉลี่ยอาจยังห่างจากศูนย์กลางมากกว่า PPO")

add_picture(doc, "ddpg_loss.png", "รูปที่ 4 Actor loss และ critic loss ของ DDPG")
add_p(doc, "critic loss ของ DDPG ต่ำในช่วงท้าย ส่วน actor loss มีค่าลบมากขึ้นตามนิยามของ DDPG ที่ optimize negative expected Q-value จึงควรอ่านร่วมกับ reward และ episode length มากกว่าตีความ loss เพียงตัวเดียว")

add_h(doc, "7. วิเคราะห์เปรียบเทียบ PPO และ DDPG")
add_p(doc, "PPO ให้ผลดีที่สุดใน run นี้ เพราะ reward ปลายทางสูงกว่าและ centered reward สูงกว่า การใช้ clipped policy update และ on-policy rollout ทำให้การเรียนรู้ไม่แกว่งมาก เหมาะกับ Unity physics environment ที่ contact และ friction ทำให้ transition dynamics noisy")
add_p(doc, "DDPG-FC-350-E-PID มีข้อดีคือโครงสร้าง off-policy actor-critic และ replay buffer ขนาดใหญ่ช่วยให้ใช้ข้อมูลซ้ำได้ อีกทั้ง architecture 2x350 ELU สอดคล้องกับงานอ้างอิง อย่างไรก็ตามผล DDPG_14 ยังมี recent reward std ประมาณ 13.50 และมี reward gap กับ PPO อยู่ จึงสรุปได้ว่า DDPG เรียนรู้การทรงตัวได้ แต่ยังต้อง tuning เพิ่มเพื่อให้เสถียรเท่า PPO")

add_h(doc, "8. ข้อจำกัดและข้อเสนอแนะ")
for item in [
    "ข้อมูลที่มีเป็น training logs เป็นหลัก ยังไม่มี evaluation log แยกที่วัดตำแหน่งลูกบอลต่อเวลา จึงยังไม่สามารถคำนวณ RMS position error หรือ overshoot จริงได้",
    "PPO และ DDPG ใช้ observation/reward ไม่เหมือนกันทั้งหมด โดย DDPG เพิ่ม E-PID terms และ progress reward ทำให้การเปรียบเทียบนี้เป็นเชิงระบบมากกว่า algorithm-only อย่างสมบูรณ์",
    "ควรเก็บ evaluation episode ด้วย seed เดียวกันสำหรับทั้งสอง model ได้แก่ PPO pid_tuning_agent-3500018.onnx และ DDPG_14/ddpg_fc_350_e_pid_agent.onnx",
    "ควรเพิ่มกราฟ trajectory x-y, radial error vs time และ PID gains vs time เพื่อยืนยันคุณภาพการควบคุมนอกเหนือจาก training reward",
]:
    add_bullet(doc, item)

add_h(doc, "9. สรุป")
add_p(
    doc,
    f"รายงานนี้อ้างอิงข้อมูลจากไฟล์โครงรายงานเดิม, source code agent, config และ TensorBoard logs ในโปรเจกต์ Unity โดยผลสุดท้ายของ PPO ใช้ pid_tuning_multi_12 และ DDPG ใช้ DDPG_14 ผลการเปรียบเทียบพบว่า PPO ให้ reward ปลายทางสูงกว่าและเสถียรกว่า โดย final reward ประมาณ {fmt(ppo_final)} ขณะที่ DDPG recent mean reward อยู่ที่ {fmt(ddpg_recent)} แต่ episode ล่าสุดของ DDPG ได้ reward {fmt(ddpg_episode_reward)} และ episode length เต็ม 2000 steps จึงแสดงว่า DDPG สามารถเรียนรู้พฤติกรรมทรงตัวได้เช่นกัน สรุปเชิงวิศวกรรมคือ PPO เป็นตัวเลือกที่พร้อมใช้และเสถียรกว่าในสถานะปัจจุบัน ส่วน DDPG-FC-350-E-PID มีศักยภาพและสอดคล้องกับงานอ้างอิง แต่ควรปรับ reward/exploration และทำ evaluation เพิ่มก่อนนำไปสรุปว่าสมรรถนะดีกว่า PPO",
)

add_h(doc, "References / เอกสารอ้างอิง")
add_p(doc, '[1] Okafor et al., "Heuristic and deep reinforcement learning-based PID control of trajectory tracking in a ball-and-plate system," 2021.')
add_p(doc, '[2] Schulman et al., "Proximal Policy Optimization Algorithms," 2017.')
add_p(doc, '[3] Lillicrap et al., "Continuous control with deep reinforcement learning," 2015.')
add_p(doc, "[4] Unity Technologies, ML-Agents Toolkit documentation and project logs from this Unity workspace.")

doc.save(OUT_DOC)
print(OUT_DOC)
